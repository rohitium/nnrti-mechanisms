#!/usr/bin/env python3
"""Apply Atanu-revision manuscript edits to DorDRM-MD-08-14-26.docx.

Writes:
  manuscript/post-feedback-from-atanu/DorDRM-MD-08-14-26-rev.docx
  manuscript/post-feedback-from-atanu/REVISION_CHANGELOG.md
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = Path("manuscript/post-feedback-from-atanu/DorDRM-MD-08-14-26.docx")
DST = Path("manuscript/post-feedback-from-atanu/DorDRM-MD-08-14-26-rev.docx")
LOG = Path("manuscript/post-feedback-from-atanu/REVISION_CHANGELOG.md")


def _set_para(p: Paragraph, text: str) -> None:
    """Replace paragraph text, keeping the first run's formatting when possible."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _clear_para(p: Paragraph) -> None:
    _set_para(p, "")


def _delete_para(p: Paragraph) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _insert_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = deepcopy(paragraph._element)
    # wipe runs in clone
    for child in list(new_p):
        if child.tag == qn("w:r"):
            new_p.remove(child)
    paragraph._element.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style is not None:
        try:
            p.style = style
        except Exception:
            pass
    _set_para(p, text)
    return p


ABSTRACT = (
    "Doravirine (DOR) is a next-generation non-nucleoside reverse transcriptase inhibitor "
    "(NNRTI) with durable clinical activity and a distinct resistance profile, but publicly "
    "available phenotypic susceptibility data remain limited for many DOR-associated genotypes. "
    "We simulated wild-type HIV-1 reverse transcriptase (RT) and 18 mutant genotypes in complex "
    "with DOR using 100 ns explicit-solvent molecular dynamics (MD; 3 replicates each) and "
    "complemented the panel with nonequilibrium free-energy perturbation (FEP) calculations "
    "that remain incomplete for several charge-change and multi-site legs. Local MD descriptors—"
    "residue–DOR contact occupancy with replicate confidence intervals, hydrogen-bond "
    "occupancy, ligand RMSF, pocket-aligned pose maps, pocket volume, and NNIBP correlated "
    "motion—identify mutation-class structural signatures. V106A-containing genotypes "
    "reproducibly increase Ser105–DOR contact occupancy; G190E strongly depletes Val179 "
    "contact; other draft-highlighted shifts (including Y188L Lys102/Pro225 and "
    "V106I+F227C–227) are descriptive under n = 3 replicate tests. Tight-SEM FEP ΔΔG values "
    "are internally precise for a subset of genotypes but correlate only weakly with "
    "experimental fold-change (Pearson R² ≈ 0.09), and V106M illustrates a binding-versus-"
    "phenotype discordance. We do not fit a phenotype-prediction model; the structural "
    "stories stand on occupancy and geometry evidence."
)

INTRO_FEP = (
    "Alchemical free-energy methods provide a complementary route to relative binding free "
    "energies (ΔΔG) for point mutations. Nonequilibrium work protocols with Crooks/BAR "
    "estimators, as implemented in pmx and related pipelines (Gapsys et al., 2020; Aldeghi et "
    "al., 2018/2019), have been used to rank NNRTI resistance mutations (Hauser et al., 2018). "
    "In this study we report an incomplete pmx/GROMACS FEP panel for DOR: protocol trust is "
    "established with a V106A walkthrough, and we discuss only genotypes whose replicate SEM "
    "is ≲ 0.6 kcal/mol. We treat weak ΔΔG–phenotype correlation and genotype-specific "
    "discordances (notably V106M) as scientific findings rather than pipeline failures, and we "
    "do not wait on unfinished charge-change legs before analyzing the MD structural suite."
)

FEP_RESULTS_HEAD = "Nonequilibrium FEP yields precise ΔΔG for a subset of genotypes but does not recover the phenotypic landscape"

FEP_RESULTS_BODY = (
    "We computed relative DOR binding free-energy changes with a dual-topology nonequilibrium "
    "protocol (hybrid WT→mutant topologies on holo and apo legs; 5 ns endpoint equilibration; "
    "100 snapshots; 500 ps forward/reverse switches; Crooks Gaussian intersection, BAR, and "
    "Jarzynski estimators; 3 replicates). A V106A protocol figure series documents the "
    "thermodynamic cycle, hybrid topology, work distributions, λ profiles where available, and "
    "estimator agreement. The completed panel is incomplete: charge-change and several "
    "compound genotypes lack reliable main-text point estimates (K103N SEM 2.19 kcal/mol; "
    "G190E charge protocol; missing K103N+M230L, K103N+L100I, and K103N+P225H). Across n = 14 "
    "completed comparisons, ΔΔG_bind versus log10(experimental fold-change) shows Pearson "
    "R² ≈ 0.09 (p ≈ 0.28); we plot the scatter without a fitted line. Genotypes with SEM ≤ "
    "0.6 kcal/mol (G190A, G190S, V106A, Y318F, Y188L, V106M) are discussed as precise relative "
    "binding estimates. V106M (+6.10 ± 0.16 kcal/mol) is tight yet far larger than its modest "
    "experimental fold (~3.4), a binding-versus-phenotype finding. MM/GBSA end-point totals "
    "from the same 100 ns trajectories remain highly replicate-noisy and are reported only as "
    "a brief negative control (Supplementary Table); they are not used for mechanistic claims."
)

SUITE_HEAD = "Local NNIBP structure and dynamics across the genotype panel"

SUITE_BODY = (
    "Beyond residue–DOR contact occupancy, we computed a pocket descriptor suite on "
    "PBC-corrected trajectories with MDAnalysis/mdtraj: DOR–protein hydrogen-bond occupancy "
    "(Lys103 backbone remains high across F227C backgrounds when mutation-site artifacts are "
    "masked), ligand RMSF after NNIBP Cα alignment, frame-level DOR pose PCA after NNIBP "
    "alignment and ligand-COM centering, ligand COM in a common WT NNIBP frame (depth r and "
    "polar angle α from the WT mean COM axis), pocket-volume proxies, NNIBP Cα motion PCA, "
    "DCCM difference maps, and residue–residue contact-frequency networks. These descriptors "
    "support the mechanism case studies below and are shown as panel-wide SI figures with "
    "story callouts in the main text."
)

OCC_STATS_SENTENCE = (
    " For every draft-highlighted Δoccupancy we report replicate mean ± 95% CI (n = 3) and a "
    "Welch test versus WT, in addition to the existing trajectory-label permutation null. "
    "Shifts that fail both the global null and the per-contact test are treated as descriptive "
    "only. Under this standard, Ser105 gains in V106A-containing genotypes and Val179 loss in "
    "G190E are supported; Y188L Lys102/Pro225 and V106I+F227C residue-227 shifts remain "
    "descriptive."
)

METHODS_FF_EXTRA = (
    " Ligand electrostatics used Gasteiger charges for reproducibility; these are not "
    "AM1-BCC/RESP and are one reason continuum electrostatic end-point estimates are not "
    "trusted. OpenFF/SMIRNOFF supplied bonded/nonbonded DOR parameters via direct chemical "
    "perception. Histidine tautomer and protonation states followed standard PDB2PQR/Amber "
    "assignment at pH 7. Production used Langevin dynamics (middle scheme, 1 ps−1 friction, "
    "2 fs) with a Monte Carlo barostat (1 bar); no separate long NPT equilibration was run "
    "beyond heating and the barostatted production segment. Convergence of local reporters "
    "(Cα RMSD, DOR–RT COM, Ser105/Val179/227 distances) is shown as coordinate-versus-time "
    "traces with running means (Supplementary); we do not apply an a-priori RMSD cutoff to "
    "discard replicates. FEP used the same ligand chemistry exported for GROMACS/pmx so that "
    "OpenMM MD and GROMACS FEP charges are stated consistently; unfinished charge-change legs "
    "are omitted from main-text point estimates."
)

METHODS_FEP = (
    "Relative binding free energies were estimated with pmx dual-topology nonequilibrium "
    "switching on GROMACS. For each leg, hybrid topologies morph WT→mutant in the DOR-bound "
    "(holo) and unbound (apo) states. After energy minimization and 5 ns endpoint "
    "equilibration, 100 snapshots were extracted and switched for 500 ps in forward and "
    "reverse directions. Free-energy changes were estimated with Crooks Gaussian intersection, "
    "BAR, and Jarzynski averaging; ΔΔG_bind = ΔG_holo − ΔG_apo (sums of additive legs for "
    "compounds). Three independent replicates per genotype were combined; confidence is "
    "reported as replicate SEM and estimator agreement, not experimental match."
)

METHODS_OCC = (
    "Highlighted contact Δoccupancy values were summarized as replicate means with Student-t "
    "95% confidence intervals (n = 3) and compared to WT with Welch’s t-test. A trajectory-"
    "label permutation null (10,000 shuffles) provided a multiple-comparison-aware global "
    "threshold for the panel-wide occupancy screen. Contacts were classified as supported "
    "(global null and Welch), mixed (exactly one), or descriptive only (neither)."
)

METHODS_SUITE = (
    "Additional pocket descriptors (hydrogen bonds, ligand RMSF, pose PCA, COM spherical "
    "maps, pocket volume, NNIBP PCA/DCCM, contact networks) were computed with MDAnalysis and "
    "mdtraj on PBC-imaged trajectories after NNIBP Cα superposition. Pose PCA centered each "
    "frame on the ligand COM after pocket alignment so that principal components reflect "
    "orientation rather than translation."
)


def main() -> int:
    doc = Document(str(SRC))
    paras = doc.paragraphs
    changelog: list[str] = []

    # --- Abstract ---
    _set_para(paras[6], "Abstract")
    _set_para(paras[7], ABSTRACT)
    changelog.append("Rewrote Abstract (no regression claim; FEP incomplete; occupancy stats).")

    # --- Intro heading + FEP lit paragraph ---
    _set_para(paras[29], "Introduction")
    # insert FEP lit after current intro block — after para 33 (before Figure 1)
    _insert_after(paras[33], INTRO_FEP)
    changelog.append("Filled Intro FEP-methods paragraph; removed placeholder heading.")

    # --- Fix Non-Nucleotide if still present elsewhere ---
    for p in doc.paragraphs:
        if "Non-Nucleotide" in p.text:
            _set_para(p, p.text.replace("Non-Nucleotide", "non-nucleoside").replace(
                "Non-nucleotide", "non-nucleoside"
            ))
    changelog.append("Fixed Non-Nucleotide → non-nucleoside.")

    # --- Table 1 K103N+M230L category ---
    t0 = doc.tables[0]
    for row in t0.rows:
        if "K103N+M230L" in row.cells[1].text:
            if not row.cells[0].text.strip():
                row.cells[0].text = "Resistant"
                changelog.append("Filled Table 1 category for K103N+M230L → Resistant.")

    # --- Demote MM/GBSA section ---
    _set_para(
        paras[43],
        "End-point MM/GBSA totals are replicate-noisy and do not track DOR phenotype",
    )
    # Scrub hybrid-topology / 5 ns if present in MM/GBSA block
    for i in range(44, 52):
        t = paras[i].text
        if "hybrid" in t.lower() or "5 ns" in t.lower() or "5ns" in t.lower():
            _set_para(
                paras[i],
                t.replace("hybrid topologies", "bound-state snapshots")
                .replace("hybrid topology", "bound-state snapshots")
                .replace("5 ns", "end-point windows")
                .replace("5ns", "end-point windows"),
            )
    changelog.append("Demoted MM/GBSA heading; scrubbed FEP language from that block.")

    # --- Insert FEP + suite sections before local contact analysis ---
    # After Figure 2 / before "Local trajectory-level analysis..."
    anchor = paras[51]  # Figure 2 caption
    p = _insert_after(anchor, FEP_RESULTS_HEAD)
    p = _insert_after(p, FEP_RESULTS_BODY)
    p = _insert_after(p, SUITE_HEAD)
    _insert_after(p, SUITE_BODY)
    changelog.append("Inserted FEP Results + NNIBP descriptor suite sections before contact analysis.")

    # --- Occupancy stats sentence into permutation paragraph ---
    # Re-resolve paragraphs after inserts
    for p in doc.paragraphs:
        if p.text.startswith("Because many residue-level contacts were evaluated"):
            _set_para(p, p.text.rstrip() + OCC_STATS_SENTENCE)
            changelog.append("Added replicate CI / Welch / descriptive-only language to occupancy screen.")
            break

    # --- Downgrade Y188L Lys102/Pro225 claim ---
    for p in doc.paragraphs:
        if "increased DOR contact occupancy at Lys102 by 0.23" in p.text:
            _set_para(
                p,
                "For example, Y188L showed an apparent increase in Lys102 contact occupancy "
                "(Δ ≈ +0.19) and decrease at Pro225 (Δ ≈ −0.13) relative to WT, whereas Y181C "
                "was near WT at these positions. With n = 3 replicates, neither Y188L shift "
                "exceeds the global permutation null nor reaches Welch p < 0.05; we therefore "
                "treat Lys102/Pro225 as descriptive reporters of Tyr188-side remodeling rather "
                "than statistically supported occupancy biomarkers.",
            )
            changelog.append("Downgraded Y188L Lys102/Pro225 occupancy claim to descriptive_only.")
            break

    # --- Hedge V106A / G190E causal language ---
    for p in doc.paragraphs:
        t = p.text
        if t.startswith("The most plausible interpretation is that replacement of valine"):
            _set_para(
                p,
                "A consistent interpretation is that replacement of valine with alanine removes "
                "hydrophobic bulk from a direct DOR-contact site and may allow the bound ligand "
                "to redistribute toward a slightly more 103–108 loop-shifted local pose "
                "ensemble. In that ensemble, contacts with Lys104 and Ser105 can arise as "
                "secondary interactions. Ser105 Δoccupancy is supported by both the global "
                "permutation null and Welch tests; Lys104 is mixed (Welch pass, global null "
                "fail). The contrast with V106I is informative: isoleucine preserves branched "
                "hydrophobic bulk at position 106 and does not show the Ser105 occupancy gain.",
            )
        elif t.startswith("The simplest interpretation is that the G190E side chain"):
            _set_para(
                p,
                "These observations are consistent with the G190E side chain introducing a "
                "local electrostatic and steric perturbation in the β9–β10 hairpin that "
                "propagates into the Val179 face of the pocket. Val179 contact loss is "
                "supported under both the global null and Welch tests; G190A does not show the "
                "same effect. We therefore treat G190E as a plausible DOR-relevant genotype "
                "deserving further phenotypic study rather than as a fully explained resistance "
                "mechanism.",
            )
        elif "therefore explains" in t or "likely causes this shift" in t or "could, therefore, explain" in t:
            _set_para(
                p,
                t.replace("therefore explains", "is consistent with")
                .replace("likely causes this shift", "is associated with this shift")
                .replace("could, therefore, explain", "may contribute to"),
            )
    changelog.append("Hedged V106A/G190E causal verbs; kept Ser105/Val179 as supported reporters.")

    # --- V106I+F227C descriptive note ---
    for p in doc.paragraphs:
        if "V106I+F227C does not display this pattern" in p.text:
            _set_para(
                p,
                p.text.rstrip()
                + " The residue-227 occupancy loss in V106I+F227C (Δ ≈ −0.26) does not pass "
                "replicate Welch or global-null thresholds and is reported as a descriptive "
                "pose/distance observation rather than a significant occupancy biomarker.",
            )
            changelog.append("Marked V106I+F227C–227 occupancy as descriptive_only.")
            break

    # --- Delete regression Results + Discussion + Methods ---
    to_clear_prefixes = (
        "A simple linear regression model trained",
        "After identifying local structural signatures across the genotype panel, we asked whether a small number",
        "Ser105-DOR minimum distance that captures",
        "Root mean squared deviation (RMSD) of DOR backbone",
        "These two features were determined using a multi-step",
        "Where  represents the standardized",
        "Figure 8: Performance of a Linear Regression",
        "This minimal model, when applied across the full panel",
        "To estimate uncertainty in model predictions, we performed bootstrap",
        "Finally, starting from a large set of MD-derived features",
        "Regression analysis of MD-derived features",
        "To test whether a compact set of trajectory-derived",
        "Uncertainty in model predictions was estimated by within-genotype bootstrap",
    )
    deleted = 0
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if any(t.startswith(pref) for pref in to_clear_prefixes):
            _delete_para(p)
            deleted += 1
        elif "custom linear regression" in t.lower() or "linear regression model" in t.lower() and "Discussion" not in t:
            # catch residual discussion sentences about the model inside longer paras already handled
            pass
    changelog.append(f"Deleted linear-regression Results/Methods/Discussion blocks ({deleted} paragraphs).")

    # Scrub any remaining regression sentences in Discussion
    for p in doc.paragraphs:
        if "linear regression" in p.text.lower() or "logistic regression" in p.text.lower():
            _set_para(
                p,
                "We deliberately do not fit a phenotype-prediction model to MD features. "
                "Two-feature regressions on ~14 genotypes are post-hoc and do not earn a "
                "Results claim; the structural stories above stand on occupancy and geometry "
                "evidence, including replicate CIs for highlighted contacts.",
            )
    changelog.append("Replaced residual regression Discussion with explicit no-model statement.")

    # --- Expand Methods FF / FEP / occupancy / suite ---
    for p in doc.paragraphs:
        if p.text.startswith("All MD simulations were performed with OpenMM"):
            _set_para(p, p.text.rstrip() + METHODS_FF_EXTRA)
            break
    # Find contact occupancy methods and append stats sentence; insert FEP/suite methods after it
    for p in doc.paragraphs:
        if p.text.startswith("Residue-level DOR contact occupancy was quantified"):
            _set_para(p, p.text.rstrip() + " " + METHODS_OCC)
            q = _insert_after(p, "Nonequilibrium FEP (pmx/GROMACS)")
            # style as body; heading as separate
            break
    # Better: insert dedicated method paragraphs before References
    for p in doc.paragraphs:
        if p.text.strip() == "References":
            prev = p  # insert before references by inserting after previous sibling — use insert before
            # Walk: insert headings+bodies just before References
            # python-docx: insert after the paragraph preceding References
            # Find previous
            break
    # Locate paragraph before References
    paras_now = doc.paragraphs
    ref_idx = next(i for i, p in enumerate(paras_now) if p.text.strip() == "References")
    anchor = paras_now[ref_idx - 1]
    h1 = _insert_after(anchor, "Nonequilibrium FEP (pmx/GROMACS)")
    b1 = _insert_after(h1, METHODS_FEP)
    h2 = _insert_after(b1, "NNIBP descriptor suite")
    _insert_after(h2, METHODS_SUITE)
    changelog.append("Expanded Methods: FF justification, FEP protocol, occupancy stats, descriptor suite.")

    # --- Discussion tone: incomplete FEP / no overclaim ---
    for p in doc.paragraphs:
        if p.text.startswith("We asked whether short explicit-solvent MD simulations"):
            _set_para(
                p,
                "We asked whether short explicit-solvent MD simulations, together with an "
                "incomplete nonequilibrium FEP panel, could provide a coherent structural "
                "framework for known and limited-data DOR resistance genotypes without claiming "
                "that 100 ns equilibration recovers binding thermodynamics or that FEP currently "
                "explains the full phenotypic landscape.",
            )
            break
    changelog.append("Discussion lead reframed around incomplete FEP + local MD claims.")

    doc.save(str(DST))

    LOG.write_text(
        "# Manuscript revision changelog\n\n"
        f"Source: `{SRC}`\n"
        f"Output: `{DST}`\n\n"
        + "\n".join(f"- {c}" for c in changelog)
        + "\n\n"
        "## Still manual / figure swaps\n\n"
        "- Drop Figure 8 (regression) from the Word gallery and renumber.\n"
        "- Insert FEP protocol panels (`results/analysis/fep_pmx/protocol/V106A/`) and "
        "no-fit scatter (`panel_ddg_vs_experiment.png`).\n"
        "- Insert occupancy CI plots (`results/analysis/occupancy_stats/plots/`).\n"
        "- Insert modern_md_suite highlights (pose PCA, H-bonds, COM r–α, DCCM).\n"
        "- Collapse Table 2 MM/GBSA to SI if desired; main text already demotes it.\n"
        "- Add Gapsys/Aldeghi/Hauser bibliography entries if not already cited numerically.\n"
    )
    print(f"Wrote {DST}")
    print(f"Wrote {LOG}")
    for c in changelog:
        print(" -", c)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
