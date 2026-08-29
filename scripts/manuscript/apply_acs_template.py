#!/usr/bin/env python3
"""Apply the ACS manuscript template's styles and page setup to the draft.

Approach
--------
Style the EXISTING document rather than pour its text into the template. The
draft carries four inline images (three EMF figures plus an equation) and two
tables; rebuilding from the template would mean re-creating those relationships
by hand, and any slip loses a figure silently. So:

1. copy the draft,
2. inject the ACS style definitions into its ``styles.xml`` (they do not exist
   there, and python-docx cannot assign a style a document has never heard of),
3. adopt the template's ``docDefaults`` (fonts) and section properties (page
   size, margins),
4. reassign paragraph styles by role.

Images, equations, tables and their relationships are never touched.

JCIM specifics (author guidelines, July 2026)
--------------------------------------------
Articles take an **unstructured** abstract; the draft's
Background/Methods/Results/Conclusion subheadings are reported by this script but
NOT removed -- collapsing them is an editorial rewrite, not a formatting change.
Keywords and a TOC graphic are required for Perspectives and Reviews, not
Articles, so ``BG_Keywords`` and ``SN_Synopsis_TOC`` stay unused.

Usage
-----
    python scripts/manuscript/apply_acs_template.py            # writes -ACS.docx
    python scripts/manuscript/apply_acs_template.py --check    # report only
"""

from __future__ import annotations

import argparse
import re
import shutil
import zipfile
from pathlib import Path

from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def qn(tag: str) -> str:
    prefix, local = tag.split(":")
    return f"{{{W}}}{local}"


#: draft paragraph role -> ACS styleId
ROLE_STYLE = {
    "title": "BATitle",
    "authors": "BBAuthorName",
    "address": "BCAuthorAddress",
    "email": "BIEmailAddress",
    "abstract": "BDAbstract",
    "body": "TAMainText",
    "figure_caption": "VAFigureCaption",
    "table_title": "VDTableTitle",
    "references": "TFReferencesSection",
}

SECTION_HEADINGS = {"Introduction", "Methods", "Results", "Discussion", "References",
                    "Abstract", "Conclusion", "Acknowledgments", "Supporting Information"}
ABSTRACT_SUBHEADS = {"Background", "Methods", "Results", "Conclusion"}


def classify(index: int, text: str, style: str, abstract_span: tuple[int, int]) -> str | None:
    """Return the role of a paragraph, or None to leave it alone."""
    t = text.strip()
    if not t:
        return None
    if index == 0:
        return "title"
    if style == "Bibliography":
        return "references"
    if re.match(r"^(Figure|Scheme)\s+\d+[:.]", t):
        return "figure_caption"
    if re.match(r"^Table\s+\d+[:.]", t):
        return "table_title"
    # Inside the abstract, "Methods"/"Results"/"Conclusion" are its subheadings,
    # not document sections -- check the span before the section-heading test.
    if abstract_span[0] <= index <= abstract_span[1]:
        return "abstract"
    if t in SECTION_HEADINGS:
        return "heading"
    if re.match(r"^[A-Z][\w .,'’-]+\d?\*?(,\s*[A-Z][\w .'’-]+\d?\*?)+$", t) and index <= 3:
        return "authors"
    if index <= 4 and re.search(r"University|Division|Department|School", t):
        return "address"
    if index <= 6 and ("@" in t or t.lower().startswith("*corresponding")):
        return "email"
    if abstract_span[0] <= index <= abstract_span[1]:
        return "abstract"
    return "body"


def find_abstract_span(paras: list[tuple[int, str]]) -> tuple[int, int]:
    start = end = -1
    for i, t in paras:
        if t.strip() == "Abstract":
            start = i + 1
        elif t.strip() == "Introduction" and start >= 0:
            end = i - 1
            break
    return (start, end)


def inject_styles(target_docx: Path, template_dotx: Path) -> list[str]:
    """Copy ACS style definitions and docDefaults into the target's styles.xml."""
    with zipfile.ZipFile(template_dotx) as z:
        tpl_styles = etree.fromstring(z.read("word/styles.xml"))
        tpl_doc = etree.fromstring(z.read("word/document.xml"))

    with zipfile.ZipFile(target_docx) as z:
        names = z.namelist()
        parts = {n: z.read(n) for n in names}

    tgt_styles = etree.fromstring(parts["word/styles.xml"])
    have = {s.get(qn("w:styleId")) for s in tgt_styles.findall(qn("w:style"))}

    # "Normal" must come across too: every ACS style is basedOn="Normal", so
    # without it they resolve against the DRAFT's Normal -- which carries no
    # explicit size and falls back to ~10pt, while the ACS Normal is 12pt. That
    # is why the first conversion rendered main text smaller than its captions.
    wanted = set(ROLE_STYLE.values()) | {"TCTableBody", "Normal"}
    added = []
    for s in tpl_styles.findall(qn("w:style")):
        sid = s.get(qn("w:styleId"))
        if sid not in wanted:
            continue
        if sid in have:                      # replace, do not skip
            for old in tgt_styles.findall(qn("w:style")):
                if old.get(qn("w:styleId")) == sid:
                    tgt_styles.remove(old)
        tgt_styles.append(s)
        added.append(sid)

    # adopt the template's font/paragraph defaults
    tpl_def = tpl_styles.find(qn("w:docDefaults"))
    tgt_def = tgt_styles.find(qn("w:docDefaults"))
    if tpl_def is not None and tgt_def is not None:
        tgt_styles.replace(tgt_def, tpl_def)

    parts["word/styles.xml"] = etree.tostring(
        tgt_styles, xml_declaration=True, encoding="UTF-8", standalone=True)

    # adopt the template's page setup (size, margins) -- keep the draft's own
    # headers/footers references, which live in the same sectPr
    tpl_sect = tpl_doc.find(f".//{qn('w:sectPr')}")
    if tpl_sect is not None:
        tgt_doc = etree.fromstring(parts["word/document.xml"])
        tgt_sect = tgt_doc.find(f".//{qn('w:sectPr')}")
        if tgt_sect is not None:
            for tag in ("w:pgSz", "w:pgMar"):
                src = tpl_sect.find(qn(tag))
                dst = tgt_sect.find(qn(tag))
                if src is not None:
                    if dst is not None:
                        tgt_sect.replace(dst, src)
                    else:
                        tgt_sect.append(src)
        parts["word/document.xml"] = etree.tostring(
            tgt_doc, xml_declaration=True, encoding="UTF-8", standalone=True)

    with zipfile.ZipFile(target_docx, "w", zipfile.ZIP_DEFLATED) as z:
        for n in names:
            z.writestr(n, parts[n])
    return added


def strip_direct_font(paragraph) -> None:
    """Remove run-level size and typeface so the paragraph style governs.

    Assigning a paragraph style does NOT clear direct character formatting, so a
    draft that hard-codes 11pt on its captions keeps rendering at 11pt no matter
    what style it is given. Only ``sz``/``szCs``/``rFonts`` are removed --
    bold, italic, underline, colour and super/subscript are meaning-bearing here
    (subscripts in the energy symbols, italics in the references) and are kept.
    """
    for run in paragraph.runs:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is None:
            continue
        for tag in ("w:sz", "w:szCs", "w:rFonts"):
            for el in rPr.findall(qn(tag)):
                rPr.remove(el)


def main() -> int:
    root = Path(__file__).resolve().parents[2]
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--src", type=Path,
                    default=root / "manuscript/post-feedback-from-atanu/DorDRM-MD-09-02-26.docx")
    ap.add_argument("--template", type=Path,
                    default=root / "manuscript/post-feedback-from-atanu/acstemplate_msw2011_mac.dotx")
    ap.add_argument("--out", type=Path, default=None)
    ap.add_argument("--check", action="store_true", help="Report the mapping; write nothing.")
    args = ap.parse_args()

    out = args.out or args.src.with_name(args.src.stem + "-ACS.docx")

    import docx
    src_doc = docx.Document(str(args.src))
    paras = [(i, p.text) for i, p in enumerate(src_doc.paragraphs)]
    span = find_abstract_span(paras)

    plan = []
    for i, p in enumerate(src_doc.paragraphs):
        role = classify(i, p.text, p.style.name, span)
        if role:
            plan.append((i, role, p.text.strip()[:70]))

    from collections import Counter
    print("Paragraph role mapping:")
    for role, n in Counter(r for _, r, _ in plan).most_common():
        style = ROLE_STYLE.get(role, "(bold main text)")
        print(f"  {role:16s} -> {style:22s} {n:4d} paragraphs")

    subheads = [t for _, r, t in plan if r == "abstract" and t in ABSTRACT_SUBHEADS]
    if subheads:
        print(f"\n  NOTE: abstract is structured ({', '.join(subheads)}).")
        print("  JCIM Articles require an UNSTRUCTURED abstract. These subheadings are")
        print("  left in place -- collapsing them is an editorial rewrite, not formatting.")

    # --- structural issues a style pass cannot fix -----------------------------
    issues = []
    embedded = [(i, t) for i, p in enumerate(src_doc.paragraphs)
                for t in [p.text.strip()]
                if re.search(r"(Figure|Table)\s+\d+\s*[:.]", t)
                and not re.match(r"^(Figure|Scheme|Table)\s+\d+[:.]", t)
                and len(t) > 200 and re.search(r"\((A|a)\)", t)]
    for i, t in embedded:
        m = re.search(r"(Figure|Table)\s+\d+", t)
        issues.append(f"paragraph {i}: {m.group(0)} caption is embedded mid-paragraph, "
                      f"so it cannot take VA_Figure_Caption. Split it into its own paragraph.")
    for k, tb in enumerate(src_doc.tables):
        head = tb.rows[0].cells[0].text.strip()
        if re.match(r"^Table\s+\d+[:.]", head):
            issues.append(f"table {k}: its title ('{head[:50]}...') is the table's own header "
                          f"row, not a preceding paragraph, so it cannot take VD_Table_Title.")
    if issues:
        print("\n  STRUCTURAL ISSUES (fix by hand in Word; not safe to automate):")
        for x in issues:
            print(f"    - {x}")

    if args.check:
        return 0

    shutil.copy2(args.src, out)
    added = inject_styles(out, args.template)
    print(f"\nInjected {len(added)} ACS styles: {', '.join(sorted(added))}")

    doc = docx.Document(str(out))
    applied = 0
    for i, p in enumerate(doc.paragraphs):
        role = classify(i, p.text, p.style.name, span)
        if role == "heading":
            p.style = doc.styles["TAMainText"]
            strip_direct_font(p)
            for r in p.runs:
                r.bold = True
            continue
        sid = ROLE_STYLE.get(role or "")
        if not sid:
            continue
        try:
            p.style = doc.styles[sid]
            strip_direct_font(p)
            applied += 1
        except KeyError:
            print(f"  WARN: style {sid} unavailable for paragraph {i}")
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    try:
                        p.style = doc.styles["TCTableBody"]
                        strip_direct_font(p)
                    except KeyError:
                        print("  WARN: TCTableBody unavailable")
    doc.save(str(out))

    with zipfile.ZipFile(out) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    print(f"Applied {applied} paragraph styles.")
    print(f"Media preserved: {len(media)} files ({', '.join(Path(m).name for m in media)})")
    print(f"Tables: {len(doc.tables)}")
    print(f"\nWrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
